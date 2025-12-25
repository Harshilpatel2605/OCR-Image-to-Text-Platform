import pdfplumber
import easyocr
import numpy as np 
import cv2
from docx import Document
import os
from docx.enum.text import WD_ALIGN_PARAGRAPH

reader = easyocr.Reader(['en'])  # GPU=False since you're CPU-only

def easyocr_with_layout(img):
    """
    Returns list of (text, bbox, conf)
    """
    results = reader.readtext(
        img,
        detail=1,
        paragraph=False
    )

    blocks = []
    for bbox, text, conf in results:
        blocks.append({
            "text": text,
            "bbox": bbox,
            "conf": conf
        })

    return blocks

def group_into_lines(blocks, y_threshold=12):
    lines = []

    for b in sorted(blocks, key=lambda x: x["bbox"][0][1]):
        y = b["bbox"][0][1]
        placed = False

        for line in lines:
            if abs(line["y"] - y) < y_threshold:
                line["words"].append(b)
                placed = True
                break

        if not placed:
            lines.append({
                "y": y,
                "words": [b]
            })

    # sort words inside lines
    for line in lines:
        line["words"] = sorted(line["words"], key=lambda x: x["bbox"][0][0])

    return lines

def detect_blocks(lines):
    blocks = []
    avg_len = sum(len(" ".join(w["text"] for w in l["words"])) for l in lines) / len(lines)

    for i, line in enumerate(lines):
        text = " ".join(w["text"] for w in line["words"]).strip()

        is_title = (
            text.isupper() and
            len(text) < avg_len * 0.6 and
            i < 3
        )

        blocks.append({
            "type": "title" if is_title else "paragraph",
            "text": text
        })

    return blocks

import re

def postprocess_text(text):
    # fix hyphen line breaks
    text = re.sub(r"(\w+)-\s+(\w+)", r"\1\2", text)

    # normalize spaces
    text = re.sub(r"\s{2,}", " ", text)

    return text.strip()


def preprocess_for_ocr(img):
    gray = cv2.cvtColor(img, cv2.COLOR_BGR2GRAY)
    return cv2.fastNlMeansDenoising(gray, h=20)


def extract_text_pdf(path):
    pages = []

    with pdfplumber.open(path) as pdf:
        for page in pdf.pages:
            text = page.extract_text()

            # Fast path: selectable text exists
            if text and len(text.strip()) > 50:
                pages.append(text)
                continue

            # Extra cheap check
            if page.extract_words():
                pages.append(text or "")
                continue

            # OCR fallback
            img = np.array(page.to_image(resolution=200).original)
            img = preprocess_for_ocr(img)

            result = reader.readtext(
                img,
                detail=0,
                paragraph=True
            )

            pages.append(" ".join(result))

            del img  # free memory

    return pages



def extract_text_docx(path):
    doc = Document(path)
    pages = []

    current_page = []

    for para in doc.paragraphs:
        if para.text.strip():
            current_page.append(para.text)

        # Simple page break heuristic
        if para.text.strip() == "":
            pages.append(" ".join(current_page))
            current_page = []

    if current_page:
        pages.append(" ".join(current_page))

    return pages


def extract_text_image(path):
    img = cv2.imread(path)
    if img is None:
        raise ValueError("Unable to read image")

    img = preprocess_for_ocr(img)

    blocks = easyocr_with_layout(img)
    lines = group_into_lines(blocks)
    structured = detect_blocks(lines)

    pages = []
    page_text = []

    for b in structured:
        page_text.append({
            "type": b["type"],
            "text": postprocess_text(b["text"])
        })

    pages.append(page_text)
    return pages


def extract_text(path):
    ext = os.path.splitext(path)[1].lower()

    if ext == ".pdf":
        return extract_text_pdf(path)

    elif ext == ".docx":
        return extract_text_docx(path)

    elif ext in [".jpg", ".jpeg", ".png", ".bmp", ".tiff", ".webp"]:
        return extract_text_image(path)

    else:
        raise ValueError("Unsupported file format")


def normalize_text(text):
    lines = text.splitlines()
    cleaned = []

    for line in lines:
        line = line.strip()
        if not line:
            cleaned.append("")
        else:
            cleaned.append(line)

    return "\n".join(cleaned)


def build_pages(raw_pages):
    pages = []

    for i, blocks in enumerate(raw_pages):
        pages.append({
            "page_no": i + 1,
            "blocks": blocks
        })

    return pages



## Export 1 : Plain Text
def export_txt(pages, output_path):
    with open(output_path, "w", encoding="utf-8") as f:
        for p in pages:
            f.write(p["text"])
            f.write("\n\n")

## Export 2 : DOCX


def export_docx(pages, output_path):
    doc = Document()

    for page in pages:
        for block in page["blocks"]:
            if block["type"] == "title":
                p = doc.add_heading(block["text"], level=1)
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:
                p = doc.add_paragraph(block["text"])
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        doc.add_page_break()

    doc.save(output_path)


## Export 3 : Searchable PDF





## MAIN

if __name__ == "__main__":
    path = "./Sample/scannedImage.jpg"

    raw_pages = extract_text_image(path)
    pages = build_pages(raw_pages)

    export_docx(pages, "./Output/scannedImage.docx")

